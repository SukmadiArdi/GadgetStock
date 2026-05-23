import os
import sys
import re

# Auto-install python-docx if not present
try:
    import docx
except ImportError:
    print("python-docx not found. Installing...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, hex_color):
    """Sets background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets cell padding (margins) in twentieths of a point (dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_inline_formatting(paragraph, text):
    """Parses simple inline markdown (**bold**, *italic*, `code`) and adds to paragraph."""
    # Split text by inline elements using a regex
    # Matches: **bold**, *italic*, `code`
    pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)')
    parts = pattern.split(text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(199, 37, 78) # dark pinkish red
        else:
            paragraph.add_run(part)

def create_styled_table(doc, headers, rows):
    """Creates a beautifully styled Word table with custom borders, padding, and headers."""
    table = doc.add_table(rows=0, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    # 1. Add Header Row
    hdr_cells = table.add_row().cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        
        # Parse potential markdown inside headers
        add_inline_formatting(p, title)
        
        # Apply bold to all header text runs
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255) # white text
            
        # Shading: Navy Blue (#001f3f)
        set_cell_background(hdr_cells[i], "001f3f")
        set_cell_margins(hdr_cells[i], top=140, bottom=140, left=180, right=180)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 2. Add Data Rows
    for r_idx, row_data in enumerate(rows):
        row_cells = table.add_row().cells
        # Alternate row background colors (zebra striping)
        bg_color = "F8F9FA" if r_idx % 2 == 1 else "FFFFFF"
        
        for c_idx, val in enumerate(row_data):
            # Ensure cell index is within limits
            if c_idx >= len(row_cells):
                break
            row_cells[c_idx].text = ""
            p = row_cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            
            add_inline_formatting(p, val)
            
            # Apply styling
            for run in p.runs:
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(30, 41, 59) # Slate Grey
                
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], top=100, bottom=100, left=150, right=150)
            row_cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Add spacing after table
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(12)

def convert_md_to_docx(md_path, docx_path):
    print(f"Reading {md_path}...")
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Styles config
    styles = doc.styles
    
    # Body Font Configuration
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Arial'
    normal_font.size = Pt(11)
    normal_font.color.rgb = RGBColor(30, 41, 59) # Slate Dark Grey
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(8)
    
    print("Converting Markdown to Docx...")
    
    in_code_block = False
    code_lines = []
    code_lang = ""
    
    table_rows = []
    
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # --- 1. CODE BLOCKS ---
        if stripped.startswith('```'):
            if in_code_block:
                # Close code block
                in_code_block = False
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.4)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(12)
                p.paragraph_format.line_spacing = 1.0
                
                # Code border/background effect using paragraph properties if possible, 
                # but simple format is standard
                code_text = "\n".join(code_lines)
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9.0)
                run.font.color.rgb = RGBColor(15, 23, 42) # near black
                
                # Reset
                code_lines = []
                code_lang = ""
            else:
                # Open code block
                in_code_block = True
                code_lang = stripped[3:].strip()
            i += 1
            continue
            
        if in_code_block:
            code_lines.append(line.rstrip('\n'))
            i += 1
            continue
            
        # --- 2. TABLES ---
        if stripped.startswith('|'):
            table_rows.append(stripped)
            i += 1
            continue
        elif table_rows:
            # Table finished, process it
            headers = []
            rows = []
            for t_idx, t_line in enumerate(table_rows):
                # Clean row
                cells = [c.strip() for c in t_line.split('|')[1:-1]]
                if t_idx == 0:
                    headers = cells
                elif t_idx == 1:
                    # Divider row (e.g. | --- | --- |), skip
                    continue
                else:
                    rows.append(cells)
            
            if headers:
                create_styled_table(doc, headers, rows)
            table_rows = []
            # Do not increment i, let current non-table line be processed
        
        # --- 3. HORIZONTAL RULES ---
        if stripped in ['---', '***', '___']:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("―" * 50)
            run.font.color.rgb = RGBColor(203, 213, 225) # light border grey
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            i += 1
            continue
            
        # --- 4. HEADINGS ---
        heading_match = re.match(r'^(#{1,6})\s+(.*)$', stripped)
        if heading_match:
            hashes, title_text = heading_match.groups()
            level = len(hashes)
            
            # Create a custom heading
            p = doc.add_paragraph()
            p.paragraph_format.keep_with_next = True
            
            run = p.add_run(title_text)
            run.bold = True
            
            if level == 1:
                run.font.size = Pt(20)
                run.font.color.rgb = RGBColor(0, 31, 63) # Deep Navy
                p.paragraph_format.space_before = Pt(24)
                p.paragraph_format.space_after = Pt(8)
            elif level == 2:
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(0, 31, 63) # Deep Navy
                p.paragraph_format.space_before = Pt(18)
                p.paragraph_format.space_after = Pt(6)
            elif level == 3:
                run.font.size = Pt(13)
                run.font.color.rgb = RGBColor(30, 58, 95) # Muted Navy
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after = Pt(4)
            elif level == 4:
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(71, 85, 105) # Slate
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(4)
            else:
                run.font.size = Pt(10.5)
                run.font.color.rgb = RGBColor(71, 85, 105)
                run.italic = True
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(4)
                
            i += 1
            continue
            
        # --- 5. LIST ITEMS (BULLET & NUMBERED) ---
        bullet_match = re.match(r'^[\-\*]\s+(.*)$', stripped)
        if bullet_match:
            content = bullet_match.group(1)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            add_inline_formatting(p, content)
            i += 1
            continue
            
        numbered_match = re.match(r'^(\d+)\.\s+(.*)$', stripped)
        if numbered_match:
            num, content = numbered_match.groups()
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            add_inline_formatting(p, content)
            i += 1
            continue
            
        # --- 6. NORMAL PARAGRAPHS ---
        if stripped == '':
            # Only add empty paragraph if previous line was not empty to avoid huge white spaces
            if i > 0 and lines[i-1].strip() != '':
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(4)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(8)
            add_inline_formatting(p, stripped)
            
        i += 1

    # Save Document
    print(f"Saving output to {docx_path}...")
    doc.save(docx_path)
    print("Word Document created successfully!")

if __name__ == '__main__':
    workspace_dir = r"c:\Users\user\My Project\Rekayasa Perangkat lunak\Gadestock"
    # Adjusting for capitalized S in directory
    if not os.path.exists(workspace_dir):
        workspace_dir = r"c:\Users\user\My Project\Rekayasa Perangkat lunak\GadgetStock"
        
    md_file = os.path.join(workspace_dir, "docs", "SKPL.md")
    docx_file = os.path.join(workspace_dir, "docs", "SKPL_GadgetStock.docx")
    
    os.makedirs(os.path.dirname(docx_file), exist_ok=True)
    convert_md_to_docx(md_file, docx_file)
